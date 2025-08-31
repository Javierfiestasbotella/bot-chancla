from docx import Document
import datetime

def guardar_pregunta_fallida(pregunta):
    ruta = "app/data/preguntas_no_resueltas.docx"
    try:
        doc = Document()
        if os.path.exists(ruta):
            doc = Document(ruta)

        doc.add_paragraph(f"{datetime.datetime.now()}: {pregunta}")
        doc.save(ruta)
    except Exception as e:
        print(f"Error al guardar pregunta: {e}")
