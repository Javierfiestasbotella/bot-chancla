import os
import re
import unicodedata
import smtplib
from email.message import EmailMessage
from datetime import datetime
from difflib import get_close_matches

from flask import Flask, render_template, request, send_file
from markupsafe import Markup

import google.generativeai as genai
from dotenv import load_dotenv
from lector_pdf import leer_todos_los_pdfs_en_fragmentos

# TF-IDF y similitud
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel

# Markdown → HTML
import markdown

# Word (pendientes)
from docx import Document

# Excel
from openpyxl import load_workbook

# =========================
# Utilidades
# =========================
def safe_text(s: str) -> str:
    """Elimina caracteres problemáticos (surrogates/emoji) para evitar errores de UTF-8."""
    if s is None:
        return ""
    return s.encode("utf-8", "ignore").decode("utf-8", "ignore")

def norm(s: str) -> str:
    """minúsculas, sin acentos/diacríticos y sin dobles espacios (para comparar)"""
    s = safe_text(s).lower()
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    return re.sub(r"\s{2,}", " ", s).strip()

def chunk_text(text, size=900, overlap=200):
    text = safe_text(text)
    if len(text) <= size:
        return [text]
    chunks, start = [], 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = end - overlap
    return chunks

# =========================
# Configuración de claves
# =========================
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("❌ No se encontró la clave GOOGLE_API_KEY en el archivo .env")

genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-1.5-pro-latest")

# =========================
# Carga de documentos (PDF + XLSX)
# =========================
def leer_todos_los_xlsx_en_fragmentos(carpeta):
    frags = []
    for root, _, files in os.walk(carpeta):
        for fn in files:
            if fn.lower().endswith(".xlsx"):
                path = os.path.join(root, fn)
                try:
                    wb = load_workbook(path, data_only=True)
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        lines = []
                        for row in ws.iter_rows(values_only=True):
                            vals = [str(v) for v in row if v is not None]
                            if vals:
                                lines.append(" | ".join(vals))
                        text = f"[DOC: {fn} - Hoja: {sheet_name}]\n" + "\n".join(lines)
                        for ch in chunk_text(text, size=1000, overlap=250):
                            frags.append(safe_text(ch))
                except Exception as e:
                    frags.append(safe_text(f"[DOC ERROR {fn}] {e}"))
    return frags

_fragmentos_pdf = leer_todos_los_pdfs_en_fragmentos("data/pdf_data")
_fragmentos_xlsx = leer_todos_los_xlsx_en_fragmentos("data/pdf_data")
fragmentos = [safe_text(f) for f in (_fragmentos_pdf + _fragmentos_xlsx)]
if not fragmentos:
    fragmentos = ["[No hay fragmentos cargados de los documentos.]"]

app = Flask(__name__)

# =========================
# Persistencia de pendientes (DOCX) + envío por email
# =========================
LEES_DIR = "lees_resp"
LEES_DOCX = os.path.join(LEES_DIR, "respuestas.docx")

def asegurar_docx():
    os.makedirs(LEES_DIR, exist_ok=True)
    if not os.path.exists(LEES_DOCX):
        doc = Document()
        doc.add_heading("Preguntas sin respuesta / con error", level=1)
        doc.add_paragraph(f"Documento creado el {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("")
        doc.save(LEES_DOCX)

def enviar_docx_por_email(path_docx: str, asunto: str = "Pendientes del asistente"):
    """
    Envía el archivo DOCX por email usando SMTP con credenciales de entorno.
    Requiere variables:
      EMAIL_HOST, EMAIL_PORT, EMAIL_USER, EMAIL_PASS, EMAIL_TO
    """
    host = os.getenv("EMAIL_HOST")
    port = int(os.getenv("EMAIL_PORT", "587"))
    user = os.getenv("EMAIL_USER")
    pwd  = os.getenv("EMAIL_PASS")
    to   = os.getenv("EMAIL_TO")

    if not all([host, port, user, pwd, to]):
        return  # si falta config, no rompemos

    if not os.path.exists(path_docx):
        return

    msg = EmailMessage()
    msg["Subject"] = asunto
    msg["From"] = user
    msg["To"] = to
    msg.set_content("Adjunto el archivo de preguntas pendientes del asistente.")

    with open(path_docx, "rb") as f:
        data = f.read()
    msg.add_attachment(
        data,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(path_docx),
    )

    with smtplib.SMTP(host, port) as server:
        server.starttls()
        server.login(user, pwd)
        server.send_message(msg)

def anotar_pendiente(pregunta: str, motivo: str, contexto_preview: str = ""):
    asegurar_docx()
    doc = Document(LEES_DOCX)
    doc.add_heading(datetime.now().strftime('%Y-%m-%d %H:%M'), level=2)
    doc.add_paragraph(f"Pregunta: {safe_text(pregunta)}")
    doc.add_paragraph(f"Motivo: {safe_text(motivo)}")
    if contexto_preview:
        doc.add_paragraph("Contexto usado (preview):")
        doc.add_paragraph(safe_text(contexto_preview[:1200]))
    doc.add_paragraph("")
    doc.save(LEES_DOCX)

    # Enviar siempre el DOCX actualizado por email
    try:
        enviar_docx_por_email(LEES_DOCX, asunto="Pendientes del asistente - La Chancla")
    except Exception:
        pass  # no interrumpir si falla el correo

# =========================
# Catálogo de vinos (desde fragmentos)
# =========================
def construir_catalogo_vinos(fragmentos):
    vinos = []
    for frag in fragmentos:
        lines = [l.strip() for l in frag.splitlines()]
        for i, line in enumerate(lines):
            if "📍" in line and ("D.O." in line or "D.O" in line or "rioja" in line.lower() or "ribeiro" in line.lower() or "tierras" in line.lower()):
                name = ""
                j = i - 1
                while j >= 0 and not name:
                    cand = lines[j].strip()
                    if cand and not cand.startswith(("📍", "🍇", "🛢")):
                        name = cand
                    j -= 1
                do = line.replace("📍", "").strip()
                uvas = crianza = nota = ""
                k = i + 1
                while k < len(lines):
                    l2 = lines[k]
                    if "📍" in l2: break
                    if l2.startswith("🍇"): uvas = l2.replace("🍇", "").strip()
                    elif l2.startswith("🛢"): crianza = l2.replace("🛢", "").strip()
                    elif l2: nota = (nota + " " + l2).strip() if nota else l2
                    k += 1
                def clean(x): return safe_text(re.sub(r"\s{2,}", " ", x))
                name, do, uvas, crianza, nota = map(clean, [name, do, uvas, crianza, nota])
                if name and len(name) <= 80:
                    vinos.append({"nombre": name, "do": do, "uvas": uvas, "crianza": crianza, "nota": nota})
    vistos, result = set(), []
    for v in vinos:
        n = v["nombre"].strip().lower()
        if n not in vistos:
            vistos.add(n); result.append(v)
    return result

CATALOGO_VINOS = construir_catalogo_vinos(fragmentos)

def buscar_vino_por_nombre(nombre, catalogo):
    nombres = [v["nombre"] for v in catalogo]
    candidatos = get_close_matches(nombre, nombres, n=1, cutoff=0.6)
    if candidatos:
        nombre_ok = candidatos[0]
        for v in catalogo:
            if v["nombre"] == nombre_ok:
                return v
    for v in catalogo:
        if nombre.lower() in v["nombre"].lower():
            return v
    return None

def markdown_vino(v):
    filas = [
        f"**{v['nombre']}**",
        f"- **D.O.**: {v['do']}" if v["do"] else "",
        f"- **Uvas**: {v['uvas']}" if v["uvas"] else "",
        f"- **Crianza**: {v['crianza']}" if v["crianza"] else "",
        f"- **Nota**: {v['nota']}" if v["nota"] else "",
    ]
    return "\n".join([x for x in filas if x])

def markdown_tabla_vinos(vinos):
    cab = "| Vino | D.O. | Uvas | Crianza |\n|---|---|---|---|"
    filas = [f"| {v['nombre']} | {v['do']} | {v['uvas']} | {v['crianza']} |" for v in vinos]
    return "\n".join([cab] + filas)

def es_pregunta_de_vinos(p):
    p = norm(p)
    claves = ["vino","vinos","tinto","blanco","rosado","espumoso","cava","rioja","ribeiro","ronda","malagueno","malagueño"]
    return any(c in p for c in claves)

# =========================
# Clasificación de tema + expansión
# =========================
TOPICS = {
    "staff": {"personal","trabajador","trabajadores","trabajadora","empleado","empleados","equipo",
              "plantilla","camarero","camareros","cocina","cocinero","cocineros","chef",
              "responsable","encargado","funciones","puestos","tareas","rrhh"},
    "uniform": {"uniforme","uniformes","vestimenta","ropa","polo","camiseta","pantalon","zapatos","calzado","zapatillas"},
    "payments": {"cobrar","cobro","cobros","pago","pagos","caja","arqueo","cuadre","cierre",
                 "tpv","datofono","datafono","efectivo","tarjeta","bizum","factura","ticket","nidex"},
    "schedule": {"horario","turno","turnos","entrada","llegada","salida","descanso","viernes"},
    "vinos": {"vino","vinos","tinto","blanco","rosado","cava","espumoso","rioja","ribeiro","ronda","malagueno","malagueño"},
}

EXPANSIONES = {
    "payments": ["cuadre de caja","cierre de caja","arqueo","formas de pago","como cobramos","facturacion"],
    "staff": ["equipo","funciones del personal","responsabilidades","organizacion del personal","plantilla"],
    "uniform": ["ropa de trabajo","normas de vestimenta"],
    "schedule": ["apertura","cierre","descansos"]
}

BEBIDAS_PRECIOS = {"refresco","cocacola","fanta","aquarius","nestea","ginger","bitter","sangria","cerveza",
                   "botella","copa","jarra","cafe","te","infusion","precio","precios","€"}

def detectar_topic(pregunta: str) -> str:
    p = norm(pregunta)
    mejor, puntos = "", 0
    for t, kws in TOPICS.items():
        score = sum(1 for k in kws if k in p)
        if score > puntos:
            mejor, puntos = t, score
    return mejor

def expand_query(pregunta: str, topic: str) -> str:
    ex = EXPANSIONES.get(topic, [])
    return (pregunta + " " + " ".join(ex)).strip() if ex else pregunta

def es_fragmento_de_precios(frag: str) -> bool:
    f = norm(frag)
    if "€" in frag or re.search(r"\d+,\d{2}\s*€", frag):
        return True
    return any(w in f for w in BEBIDAS_PRECIOS)

# =========================
# Búsqueda con penalizaciones/bonos por tema
# =========================
def encontrar_fragmentos_relacionados(pregunta, fragmentos, max_resultados=8):
    topic = detectar_topic(pregunta)
    q = expand_query(pregunta, topic)

    # TF-IDF palabras (1-2gram) y caracteres (3-5) para robustez
    v_words = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
    v_char = TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5), min_df=1)

    tfidf_w = v_words.fit_transform(fragmentos + [q])
    tfidf_c = v_char.fit_transform(fragmentos + [q])

    sims_w = linear_kernel(tfidf_w[-1], tfidf_w[:-1]).flatten()
    sims_c = linear_kernel(tfidf_c[-1], tfidf_c[:-1]).flatten()
    sims = 0.6 * sims_w + 0.4 * sims_c

    # Ajustes por tema
    ajustadas = []
    for i, frag in enumerate(fragmentos):
        score = sims[i]
        low = norm(frag)
        if topic in {"staff", "uniform", "payments"} and es_fragmento_de_precios(frag):
            score -= 0.15
        for k in TOPICS.get(topic, []):
            if k in low:
                score += 0.05
        ajustadas.append((score, i))

    ajustadas.sort(key=lambda x: x[0], reverse=True)
    top = ajustadas[:max_resultados]

    if not top or top[0][0] < 0.03:
        return [], topic
    return [fragmentos[i] for _, i in top], topic

# =========================
# Rutas
# =========================
@app.route("/")
def home():
    return render_template("index.html")

@app.route("/preguntar", methods=["POST"])
def preguntar():
    pregunta = safe_text(request.form.get("pregunta", "").strip())

    try:
        # --- Rama especial: Vinos ---
        if es_pregunta_de_vinos(pregunta) and CATALOGO_VINOS:
            tokens = [t for t in re.findall(r"[A-Za-zÁÉÍÓÚÑÜ][\wÁÉÍÓÚÑÜ-]*", pregunta)]
            candidato_nombre = " ".join([t for t in tokens if len(t) > 2])
            vino_encontrado = buscar_vino_por_nombre(candidato_nombre, CATALOGO_VINOS) if candidato_nombre else None
            if vino_encontrado:
                md = f"### Ficha del vino\n\n{markdown_vino(vino_encontrado)}"
            else:
                md = "### Carta de vinos (resumen)\n\n" + markdown_tabla_vinos(CATALOGO_VINOS)
            html = markdown.markdown(safe_text(md), extensions=["extra"])
            return render_template("index.html", pregunta=pregunta, respuesta=Markup(html))

        # --- Rama general (RAG) ---
        top_fragmentos, topic = encontrar_fragmentos_relacionados(pregunta, fragmentos, max_resultados=10)
        if not top_fragmentos:
            anotar_pendiente(pregunta, "No encontrado (sin fragmentos útiles)")
            html = markdown.markdown("**No encontrado en los documentos.** Ya lo he anotado para añadir la información.")
            return render_template("index.html", pregunta=pregunta, respuesta=Markup(html))

        contexto = safe_text("\n\n---\n\n".join(top_fragmentos))

        prompt = f"""Usa SOLO la información del contexto para responder.
Si el dato no aparece, responde EXACTAMENTE: NO_ENCONTRADO
Responde en **Markdown** (títulos, listas o tablas cuando ayuden). Sé breve y claro en español.

Contexto:
{contexto}

Pregunta:
{pregunta}
"""
        respuesta = model.generate_content(prompt)
        texto = safe_text((respuesta.text or "").strip())

        if texto.strip() == "NO_ENCONTRADO" or not texto:
            anotar_pendiente(pregunta, "No encontrado (modelo)", contexto_preview=contexto)
            html = markdown.markdown("**No encontrado en los documentos.** Ya lo he anotado para añadir la información.")
        else:
            html = markdown.markdown(texto, extensions=["extra"])

        return render_template("index.html", pregunta=pregunta, respuesta=Markup(html))

    except Exception as e:
        anotar_pendiente(pregunta, f"Error: {safe_text(str(e))}")
        html = markdown.markdown(f"Error al generar respuesta: {safe_text(str(e))}")
        return render_template("index.html", pregunta=pregunta, respuesta=Markup(html))

# Descarga del DOCX desde la web
@app.route("/descargar_pendientes")
def descargar_pendientes():
    path = os.path.join(LEES_DIR, "respuestas.docx")
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    else:
        return "Aún no hay archivo de pendientes."

# Salud y diagnóstico
@app.route("/health")
def health():
    return "ok"

@app.route("/_routes")
def list_routes():
    output = []
    for rule in app.url_map.iter_rules():
        methods = ",".join(sorted(m for m in rule.methods if m not in ("HEAD","OPTIONS")))
        output.append(f"{methods:10s} {rule.rule}")
    return "<pre>" + "\n".join(sorted(output)) + "</pre>"

# =========================
# Arranque (local/Render)
# =========================
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
